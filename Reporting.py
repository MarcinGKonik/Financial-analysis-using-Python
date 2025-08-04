import yfinance as yf
#yf.enable_debug_mode()
import pandas as pd 
import os
import openai 
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from langchain_community.tools.yahoo_finance_news import YahooFinanceNewsTool
from langgraph.prebuilt import create_react_agent

#os.environ["OPENAI_API_KEY"] = "YOUR_API_KEY"

#Downloading stock data from yfinance
ticker = "AAPL" 
print(f"Downloading historic data on {ticker}")
stock = yf.Ticker(ticker)
hist = stock.history(period="30d")
print(hist[['Open', 'Close']])

#Visualizing stock data
plt.figure(figsize=(10, 5))
plt.plot(hist.index, hist['Close'], label='Close Price', color='blue', linewidth=2)
plt.title(f"{ticker} - graph")
plt.xlabel("Date")
plt.ylabel("Price")
plt.legend()
plt.grid(True)
print("Visual data plot generated")

#Saving visualization plot
chart_path = "price_chart.png"
plt.savefig(chart_path)
plt.close()
print("Visual data plot saved")


#Downloading latest news Headline
# tool = [YahooFinanceNewsTool()]
# agent = create_react_agent("openai:gpt-4.1-mini", tool)

# input_message = {
#     "role": "user",
#     "content": f"Today's news related to {ticker} stock's on yahoo finance website"
# }

# response_text = ""

# for step in agent.stream(
#     {"messages": [input_message]},
#     stream_mode="values",
# ):
#     step["messages"][-1].pretty_print()
#     response_text += step["messages"][-1].content

#Creating word report
if ticker:
    doc = Document()
    doc.add_heading(f'Report for {ticker}', 0)

    doc.add_heading('Latest 10 Days - Open & Close prices', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Open Price'
    hdr_cells[2].text = 'Close Price'

    for date, row in hist.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(date.date())
        row_cells[1].text = f"{row['Open']:.2f}"
        row_cells[2].text = f"{row['Close']:.2f}"
    
    #adding data visualization to the report using docx Inches
    doc.add_heading('Price Chart (Last 30 days)', level=1)
    doc.add_picture(chart_path, width=Inches(6))
    print("Visualization added to the report")

    #adding news to the report
    # oc.add_heading('recent news headlines', level=1)
    # doc.add_paragraph(response_text)
     
    doc.save("report.docx")
    print("Report.docx generated")
    #TODO rewrite hist so its compatible with excel
    #export to excel
    hist.index = hist.index.tz_localize(None)
    hist[['Open', 'Close']].to_excel("reportExcel.xlsx")
    print("Report.xlsx generated")
else: 
    print("Couldn't export docs due to missing ticker")



